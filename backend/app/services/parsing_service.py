import os
import logging
from langchain_core.documents import Document as LCDocument

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "txt"}


def get_file_type(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in SUPPORTED_TYPES:
        return ext
    raise ValueError(f"Unsupported file type: {ext}")


def parse_document(file_path: str, file_type: str) -> list[LCDocument]:
    """Parse a document file into LangChain Document objects."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info(f"Parsing {file_type} file: {file_path}")

    if file_type == "pdf":
        return _parse_pdf(file_path)
    elif file_type == "docx":
        return _parse_docx(file_path)
    elif file_type == "txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _parse_pdf(file_path: str) -> list[LCDocument]:
    import fitz  # pymupdf

    docs = []
    pdf = fitz.open(file_path)
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text()
        if text.strip():
            docs.append(
                LCDocument(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "page_number": page_num + 1,
                        "total_pages": len(pdf),
                    },
                )
            )
    pdf.close()
    logger.info(f"Parsed {len(docs)} pages from PDF")
    return docs


def _parse_docx(file_path: str) -> list[LCDocument]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    text = "\n".join(full_text)
    if text.strip():
        return [
            LCDocument(
                page_content=text,
                metadata={"source": file_path, "page_number": 1},
            )
        ]
    return []


def _parse_txt(file_path: str) -> list[LCDocument]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    if text.strip():
        return [
            LCDocument(
                page_content=text,
                metadata={"source": file_path, "page_number": 1},
            )
        ]
    return []
